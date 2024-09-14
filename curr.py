from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.size = Pt(12)

def add_paragraph(doc, label, text):
    paragraph = doc.add_paragraph()
    run_label = paragraph.add_run(label)
    run_label.bold = True
    run_label.font.size = Pt(10)
    
    run_text = paragraph.add_run(' ' + text)
    run_text.font.size = Pt(10)
    
    return paragraph

def generate_european_cv():
    doc = Document()
    title = doc.add_heading('Curriculum Vitae Europass', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, 'Informazioni personali', level=1)
    add_paragraph(doc, 'Nome:', 'Mario Rossi')
    add_paragraph(doc, 'Indirizzo:', 'Via Roma 1, 00100 Roma, Italia')
    add_paragraph(doc, 'Telefono:', '+39 06 1234567')
    add_paragraph(doc, 'E-mail:', 'mario.rossi@example.com')
    add_paragraph(doc, 'Nazionalità:', 'Italiana')
    add_paragraph(doc, 'Data di nascita:', '01/01/1990')
    add_paragraph(doc, 'Sesso:', 'Maschile')

    add_heading(doc, 'Esperienza lavorativa', level=1)
    add_paragraph(doc, 'Periodo:', 'Gennaio 2020 - Presente')
    add_paragraph(doc, 'Datore di lavoro:', 'Azienda ABC S.p.A.')
    add_paragraph(doc, 'Posizione ricoperta:', 'Ingegnere del software')
    add_paragraph(doc, 'Principali attività e responsabilità:', 
                  'Sviluppo di applicazioni web, gestione del team di sviluppo, coordinamento con altri reparti aziendali.')

    add_heading(doc, 'Istruzione e formazione', level=1)
    add_paragraph(doc, 'Periodo:', 'Settembre 2010 - Luglio 2015')
    add_paragraph(doc, 'Titolo della qualifica rilasciata:', 'Laurea in Ingegneria Informatica')
    add_paragraph(doc, 'Istituto:', 'Università degli Studi di Roma "La Sapienza"')
    add_paragraph(doc, 'Principali materie / competenze professionali apprese:', 
                  'Programmazione, algoritmi, database, reti di calcolatori.')

    add_heading(doc, 'Competenze personali', level=1)
    add_paragraph(doc, 'Lingua madre:', 'Italiano')
    add_paragraph(doc, 'Altre lingue:', 'Inglese - Livello B2')
    add_paragraph(doc, 'Competenze comunicative:', 
                  'Buone capacità di comunicazione sviluppate grazie all’esperienza di lavoro in team.')
    add_paragraph(doc, 'Competenze organizzative:', 
                  'Esperienza nella gestione di progetti e coordinamento di piccoli gruppi di lavoro.')
    add_paragraph(doc, 'Competenze informatiche:', 
                  'Conoscenza avanzata di Python, Java, HTML/CSS, database SQL.')

    add_heading(doc, 'Ulteriori informazioni', level=1)
    add_paragraph(doc, 'Pubblicazioni:', 'Articolo scientifico sulla rivista X.')
    add_paragraph(doc, 'Conferenze:', 'Partecipazione alla conferenza Y.')
    add_paragraph(doc, 'Referenze:', 'Disponibili su richiesta.')

    doc.save('curriculum_vitae_europeo.docx')
    print("Curriculum Vitae salvato come 'curriculum_vitae_europeo.docx'")

generate_european_cv()
